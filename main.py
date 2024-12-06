import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt
from xml.etree.ElementTree import Element as OxmlElement
from xml.etree.ElementTree import QName as qn
from tkcalendar import DateEntry
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


def insertar_tabla(doc, paragraph, horarios):
    try:
        table = doc.add_table(rows=1, cols=4)
        headers = ['Tipo de Horario', 'Turno', 'Horario', 'Días']        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        
        for horario in horarios:
            row_cells = table.add_row().cells
            row_data = [horario['tipo'], horario['turno'], horario['horario'], horario['dias']]
            for i, data in enumerate(row_data):
                row_cells[i].text = data
                run = row_cells[i].paragraphs[0].runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(11)

        tbl = table._tbl
        paragraph._element.addnext(tbl)

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)
    except Exception as e:
        print(f"Error al insertar tabla: {e}")
    
def capturar_seleccionados(roles):
    try:
        return [rol for rol, var in roles.items() if var.get()]
    except Exception as e:
        print(f"Error al capturar sanciones: {e}")
        return []

def reemplazar_datos_en_plantilla(datos):
    try:
        doc = Document('template.docx')
        keys_to_uppercase = ['NOMBRE', 'MUNICIPIO', 'DEPARTAMENTO', 'REPRESENTANTE_LEGAL']  

        for p in doc.paragraphs:
            for run in p.runs:
                for key, value in datos.items():
                    if key in keys_to_uppercase:
                        value = value.upper()  # Convertir a mayúsculas
                    if f"|{key}|" in run.text:
                        print(f"Encontrado marcador |{key}| en el párrafo: {run.text}")
                        run.text = run.text.replace(f"|{key}|", value)
                        if key == 'NOMBRE':
                            run.bold = True
        for p in doc.paragraphs:
            if "|HORARIO|" in p.text:
                p.text = p.text.replace('|HORARIO|', "")
                insertar_tabla(doc, p, datos['horarios'])
                break
        doc.save('documento_completado.docx')
        print("Documento generado correctamente")
    except Exception as e:
        print(f"Error al reemplazar datos en plantilla: {e}")

def agregar_fila(tipo, row, frame):
    try:
        font_settings = ("Helvetica", 14)
        entry_width = 20

        tk.Label(frame, text=tipo, font=font_settings).grid(row=row, column=0, sticky="w")
        
        entry_turno = tk.Entry(frame, font=font_settings, width=entry_width)
        entry_turno.grid(row=row, column=1)
        
        entry_horario = tk.Entry(frame, font=font_settings, width=entry_width)
        entry_horario.grid(row=row, column=2)
        
        entry_dias = tk.Entry(frame, font=font_settings, width=entry_width)
        entry_dias.grid(row=row, column=3)
        
        entry_widgets.append({"tipo": tipo, "entry_turno": entry_turno, "entry_horario": entry_horario, "entry_dias": entry_dias})
    except Exception as e:
        print(f"Error al agregar fila: {e}")

def agregar_filas(tipo, cantidad, frame, start_row):
    for i in range(cantidad):
        agregar_fila(tipo, start_row + i + 1, frame)

def generar_tabla():
    try:
        for widget in admin_frame.winfo_children():
            widget.destroy()
        for widget in oper_frame.winfo_children():
            widget.destroy()

        entry_widgets.clear()

        row = 1
        font_settings = ("Helvetica", 14)

        tk.Label(admin_frame, text="Tipo de Horario", font=font_settings).grid(row=row, column=0, sticky="nw")
        tk.Label(admin_frame, text="Turno", font=font_settings).grid(row=row, column=1, sticky="nw")
        tk.Label(admin_frame, text="Horario", font=font_settings).grid(row=row, column=2, sticky="nw")
        tk.Label(admin_frame, text="Días", font=font_settings).grid(row=row, column=3, sticky="nw")

        row += 1

        tk.Label(oper_frame, text="Tipo de Horario", font=font_settings).grid(row=row, column=0, sticky="nw")
        tk.Label(oper_frame, text="Turno", font=font_settings).grid(row=row, column=1, sticky="nw")
        tk.Label(oper_frame, text="Horario", font=font_settings).grid(row=row, column=2, sticky="nw")
        tk.Label(oper_frame, text="Días", font=font_settings).grid(row=row, column=3, sticky="nw")

        cantidad_administrativos = int(administrativo_cb.get() or 0)
        cantidad_operativos = int(operativo_cb.get() or 0)

        if cantidad_administrativos > 0:
            agregar_filas("Administrativo", cantidad_administrativos, admin_frame, 1)
        if cantidad_operativos > 0:
            agregar_filas("Operativo", cantidad_operativos, oper_frame, 1)

        canvas.update_idletasks()
    except Exception as e:
        print(f"Error al generar tabla: {e}")

def validar_campos():
    try:
        if not nombre_entry.get() or not municipio_entry.get() or not departamento_entry.get() or not fecha_pago_entry.get() or not objeto_social_entry.get():
            messagebox.showwarning("Campos Vacíos", "Por favor, complete todos los campos del formulario.")
            return False
        return True
    except Exception as e:
        print(f"Error al validar campos: {e}")
        return False

def on_submit():        
    if validar_campos():
        try:
            datos = {
                'NOMBRE': nombre_entry.get(),
                'MUNICIPIO': municipio_entry.get(),
                'DEPARTAMENTO': departamento_entry.get(),
                'OBJETO_SOCIAL': objeto_social_entry.get(),
                'FECHA_PAGO': fecha_pago_entry.get(),
                'MUNICIPIO1': municipio1_entry.get(),
                'FECHA1': fecha1_entry.get(),
                'FECHA2': fecha2_entry.get(),
                'REPRESENTANTE_LEGAL': representante_legal_entry.get(),
                'horarios': []
            }

            datos['horarios'] = [
                {
                    "tipo": item["tipo"],
                    "turno": item["entry_turno"].get(),
                    "horario": item["entry_horario"].get(),
                    "dias": item["entry_dias"].get()
                } for item in entry_widgets
            ]

            datos['ORDEN_JERARQUICO'] = "\n".join(capturar_seleccionados(orden_jerarquico_vars))
            datos['IMPONER_SANCIONES'] = "\n".join(capturar_seleccionados(imponer_sanciones_vars))

            # Imprimir el contenido de imponer_sanciones_vars para depuración
            print("imponer_sanciones_vars:", imponer_sanciones_vars)
            print("IMPONER_SANCIONES:", datos['IMPONER_SANCIONES'])

            reemplazar_datos_en_plantilla(datos)
            messagebox.showinfo("Éxito", "El documento se ha generado correctamente.")
        except Exception as e:
            print(f"Error en on_submit: {e}")

def aceptar():
    try:
        cantidad_administrativos = int(administrativo_cb.get() or 0)
        cantidad_operativos = int(operativo_cb.get() or 0)

        if cantidad_administrativos > 0 or cantidad_operativos > 0:
            generar_tabla()
            operativo_cb.config(state=tk.DISABLED)
            administrativo_cb.config(state=tk.DISABLED)
        else:
            for widget in admin_frame.winfo_children():
                widget.destroy()
            for widget in oper_frame.winfo_children():
                widget.destroy()
            entry_widgets.clear()
            operativo_cb.config(state=tk.NORMAL)
            administrativo_cb.config(state=tk.NORMAL)
    except Exception as e:
        print(f"Error en aceptar: {e}")

def toggle_operativo_cb():
    operativo_cb.config(state=tk.NORMAL if operativo_var.get() else tk.DISABLED)

def toggle_administrativo_cb():
    administrativo_cb.config(state=tk.NORMAL if administrativo_var.get() else tk.DISABLED)

# Variables globales
bg_color = '#b0d4ec'
entry_widgets = []
font_style = ("Helvetica", 14, "italic")

# Definición de la interfaz gráfica
ventana = tk.Tk()
ventana.title("Formulario de Datos")
ventana.configure(bg=bg_color)

# Variables para los checkbuttons
gerente_var = tk.IntVar()
subgerente_var = tk.IntVar()
lider_talento_humano_var = tk.IntVar()
coordinador_sistemas_var = tk.IntVar()
lider_operativo_var = tk.IntVar()
supervisores_var = tk.IntVar()
operarios_manual_var = tk.IntVar()
operativo_var = tk.IntVar()
administrativo_var = tk.IntVar()

# Diccionario para las variables de los checkbuttons
imponer_sanciones_vars = {
    "Gerente": gerente_var,
    "Subgerente": subgerente_var,
    "Líder de talento humano": lider_talento_humano_var,    
    "Supervisores": supervisores_var,    
}

orden_jerarquico_vars = {
    "Gerente": gerente_var,
    "Subgerente": subgerente_var,
    "Líder de talento humano": lider_talento_humano_var,
    "Coordinador de sistemas integrados de gestión": coordinador_sistemas_var,
    "Líder Operativo": lider_operativo_var,
    "Supervisores": supervisores_var,
    "Operarios manual": operarios_manual_var
}


# Crear un canvas y un frame para el contenido
canvas = tk.Canvas(ventana, bg=bg_color)
scroll_y = tk.Scrollbar(ventana, orient="vertical", command=canvas.yview)
scroll_x = tk.Scrollbar(ventana, orient="horizontal", command=canvas.xview)

# Frame que contendrá todos los widgets
frame_contenido = tk.Frame(canvas, bg=bg_color)

# Configurar el canvas
canvas.create_window((0, 0), window=frame_contenido, anchor="nw")
canvas.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

# Empaquetar el canvas y la scrollbar
canvas.grid(row=0, column=0, sticky="nsew")

scroll_y.grid(row=0, column=1, sticky="ns")
scroll_x.grid(row=1, column=0, sticky="ew")

# Configurar la expansión del canvas
ventana.grid_rowconfigure(0, weight=1)
ventana.grid_columnconfigure(0, weight=1)

# Actualizar el tamaño del canvas para que se ajuste al contenido
frame_contenido.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

# Frame para los datos personales (Nombre, Municipio, Departamento)
frame_datos = tk.Frame(frame_contenido, bg=bg_color)
frame_datos.grid(padx=10, pady=10, sticky="nsew")

for i in range(6):
    frame_datos.columnconfigure(i, weight=1)

# Nombre Empresa
tk.Label(frame_datos, text="Nombre Empresa:", font=font_style, bg=bg_color).grid(row=0, column=0, sticky="e")
nombre_entry = tk.Entry(frame_datos, font=font_style)
nombre_entry.grid(row=0, column=1, padx=(0, 20), sticky="ew")

# Departamento
tk.Label(frame_datos, text="Departamento:", font=font_style, bg=bg_color).grid(row=0, column=2, sticky="e")
departamento_entry = tk.Entry(frame_datos, font=font_style)
departamento_entry.grid(row=0, column=3, padx=(0, 20), sticky="ew")

# Municipio
tk.Label(frame_datos, text="Municipio:", font=font_style, bg=bg_color).grid(row=0, column=4, sticky="e")
municipio_entry = tk.Entry(frame_datos, font=font_style)
municipio_entry.grid(row=0, column=5, pady=20, sticky="ew")

# Frame para el subtítulo y las cajas de texto (CAPITULO V)
frame_capitulo_v = tk.Frame(frame_contenido, bg=bg_color)
frame_capitulo_v.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")

tk.Label(frame_capitulo_v, text="CAPITULO V", font=font_style, bg=bg_color).grid(row=0, column=0, columnspan=2, sticky="w")

tk.Label(frame_capitulo_v, text="MUNICIPIO1:", font=font_style, bg=bg_color).grid(row=1, column=0, sticky="e")
municipio1_entry = tk.Entry(frame_capitulo_v, font=font_style)
municipio1_entry.grid(row=1, column=1, pady=20, sticky="ew")

# Objeto Social
tk.Label(frame_datos, text="Objeto Social:", font=font_style, bg=bg_color).grid(row=2, column=0, sticky="e")
objeto_social_entry = tk.Entry(frame_datos, font=font_style)
objeto_social_entry.grid(row=2, column=1, columnspan=5, sticky="ew", pady=20)

# Fecha de Pago
tk.Label(frame_datos, text="Fecha de Pago:", font=font_style, bg=bg_color).grid(row=3, column=0, sticky="e")
opciones_pago = ["los días 30 de cada mes", "los días 15 y 30 de cada mes", "catorcenales", "semanales"]
fecha_pago_entry = ttk.Combobox(frame_datos, values=opciones_pago, font=font_style, state="readonly")
fecha_pago_entry.grid(row=3, column=1, columnspan=5, sticky="ew", pady=20)
fecha_pago_entry.set("Seleccione una opción")

# Frame para los checkbuttons (Horario de trabajo)
frame_horarios = tk.Frame(frame_contenido, bg=bg_color)
frame_horarios.grid(row=4, column=0, sticky="nsew")

# Configurar las filas para que sean responsivas
frame_horarios.rowconfigure(0, weight=1)

tk.Label(frame_horarios, text="Horario de trabajo:", bg=bg_color, font=font_style).grid(row=0, column=0, sticky="w")

tk.Checkbutton(frame_horarios, text="Horario de trabajo personal operativo", variable=operativo_var, command=toggle_operativo_cb, bg=bg_color, font=("Helvetica", 14, "italic")).grid(row=1, column=0, sticky="w")
operativo_cb = tk.Entry(frame_horarios, width=5, state=tk.DISABLED)  # Inicialmente deshabilitado
operativo_cb.grid(row=1, column=1, sticky="w")

tk.Checkbutton(frame_horarios, text="Horario de trabajo personal administrativo", variable=administrativo_var, command=toggle_administrativo_cb, bg=bg_color, font=("Helvetica", 14, "italic")).grid(row=2, column=0, sticky="w")
administrativo_cb = tk.Entry(frame_horarios, width=5, state=tk.DISABLED)  # Inicialmente deshabilitado
administrativo_cb.grid(row=2, column=1, sticky="w")

# Frame para la tabla
table_frame = tk.Frame(frame_contenido)
table_frame.grid(row=3, column=0, padx=10, pady=10, sticky="nsew")

# Crear frames para administrativos y operativos
admin_frame = tk.Frame(table_frame)
admin_frame.grid(row=0, column=0, padx=5, pady=5)

oper_frame = tk.Frame(table_frame)
oper_frame.grid(row=1, column=0, padx=5, pady=5)

# Botón para aceptar
btn_aceptar = tk.Button(frame_horarios, text="Aceptar", command=aceptar, bg=bg_color, font=("Helvetica", 14))
btn_aceptar.grid(row=3, column=0, columnspan=2, pady=10)

# Configurar la fuente del menú desplegable
ventana.option_add('*TCombobox*Listbox.font', font_style)

# Frame para los checkbuttons (Orden Jerárquico)
frame_orden_jerarquico = tk.Frame(frame_contenido, bg=bg_color)
frame_orden_jerarquico.grid(padx=10, pady=10, sticky="nsew")

tk.Label(frame_orden_jerarquico, text="Orden Jerárquico:", bg=bg_color, font=font_style).grid(row=0, column=0, sticky="w")

# Crear los checkbuttons
gerente_cb = tk.Checkbutton(frame_orden_jerarquico, text="Gerente", variable=gerente_var, bg=bg_color, font=font_style)
gerente_cb.grid(row=1, column=0, sticky="w")

subgerente_cb = tk.Checkbutton(frame_orden_jerarquico, text="Subgerente", variable=subgerente_var, bg=bg_color, font=font_style)
subgerente_cb.grid(row=2, column=0, sticky="w")

lider_talento_humano_cb = tk.Checkbutton(frame_orden_jerarquico, text="Líder de talento humano", variable=lider_talento_humano_var, bg=bg_color, font=font_style)
lider_talento_humano_cb.grid(row=3, column=0, sticky="w")

coordinador_sistemas_cb = tk.Checkbutton(frame_orden_jerarquico, text="Coordinador de sistemas integrados de gestión", variable=coordinador_sistemas_var, bg=bg_color, font=font_style)
coordinador_sistemas_cb.grid(row=4, column=0, sticky="w")

lider_operativo_cb = tk.Checkbutton(frame_orden_jerarquico, text="Líder Operativo", variable=lider_operativo_var, bg=bg_color, font=font_style)
lider_operativo_cb.grid(row=5, column=0, sticky="w")

supervisores_cb = tk.Checkbutton(frame_orden_jerarquico, text="Supervisores", variable=supervisores_var, bg=bg_color, font=font_style)
supervisores_cb.grid(row=6, column=0, sticky="w")

operarios_manual_cb = tk.Checkbutton(frame_orden_jerarquico, text="Operarios manual", variable=operarios_manual_var, bg=bg_color, font=font_style)
operarios_manual_cb.grid(row=7, column=0, sticky="w")

# Frame para los checkbuttons (Imponer sanciones)
frame_imponer_sanciones = tk.Frame(frame_contenido, bg=bg_color)
frame_imponer_sanciones.grid(padx=10, pady=10, sticky="nsew")


tk.Label(frame_imponer_sanciones, text="Imponer sanciones:", bg=bg_color, font=font_style).grid(row=0, column=0, sticky="w")

# Crear los checkbuttons
gerente_cb = tk.Checkbutton(frame_imponer_sanciones, text="Gerente", variable=gerente_var, bg=bg_color, font=font_style)
gerente_cb.grid(row=1, column=0, sticky="w")

subgerente_cb = tk.Checkbutton(frame_imponer_sanciones, text="Subgerente", variable=subgerente_var, bg=bg_color, font=font_style)
subgerente_cb.grid(row=2, column=0, sticky="w")

lider_talento_humano_cb = tk.Checkbutton(frame_imponer_sanciones, text="Líder de talento humano", variable=lider_talento_humano_var, bg=bg_color, font=font_style)
lider_talento_humano_cb.grid(row=3, column=0, sticky="w")

supervisores_cb = tk.Checkbutton(frame_imponer_sanciones, text="Supervisores", variable=supervisores_var, bg=bg_color, font=font_style)
supervisores_cb.grid(row=6, column=0, sticky="w")

# Frame para el subtítulo y las cajas de texto (PUBLICACIONES Y VIGENCIA)
frame_publicaciones_vigencia = tk.Frame(frame_contenido, bg=bg_color)
frame_publicaciones_vigencia.grid(row=9, column=0, padx=10, pady=10, sticky="nsew")

tk.Label(frame_publicaciones_vigencia, text="PUBLICACIONES Y VIGENCIA", font=font_style, bg=bg_color).grid(row=0, column=0, columnspan=2, sticky="w")

tk.Label(frame_publicaciones_vigencia, text="FECHA1:", font=font_style, bg=bg_color).grid(row=1, column=0, sticky="e")
fecha1_entry = DateEntry(frame_publicaciones_vigencia, font=font_style, date_pattern='yyyy-mm-dd')
fecha1_entry.grid(row=1, column=1, pady=20, sticky="ew")

tk.Label(frame_publicaciones_vigencia, text="FECHA2:", font=font_style, bg=bg_color).grid(row=2, column=0, sticky="e")
fecha2_entry = DateEntry(frame_publicaciones_vigencia, font=font_style, date_pattern='yyyy-mm-dd')
fecha2_entry.grid(row=2, column=1, pady=20, sticky="ew")

# Representante Legal
tk.Label(frame_contenido, text="Representante Legal:", font=font_style, bg=bg_color).grid(row=10, column=0, sticky="e")
representante_legal_entry = tk.Entry(frame_contenido, font=font_style)
representante_legal_entry.grid(row=10, column=1, padx=10, pady=10, sticky="ew")

# Frame para el botón de enviar
frame_botones = tk.Frame(frame_contenido, bg=bg_color)
frame_botones.grid(row=11, column=0, padx=10, pady=10, sticky="ew")


# Botón para enviar el formulario
submit_button = tk.Button(frame_botones, text="Generar Documento", command=on_submit)
submit_button.grid(row=0, column=0, pady=10, padx=10)

# Aplicar estilos al botón
submit_button.config(bg="blue", fg="white", font=("Helvetica", 12, "bold"))

# Iniciar el bucle de la aplicación Tkinter
ventana.mainloop()


